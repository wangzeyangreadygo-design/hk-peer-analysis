#!/usr/bin/env python3
"""
Build an HK peer-analysis Word document (.docx) from a filled report JSON.

Input schema matches references/report_template.md sections.
Output: .docx using 宋体 (SimSun) for body, 黑体 (SimHei) for section headings.

Usage:
  python build_docx.py report_filled.json --out Peer_Analysis_2025H1.docx
"""
import argparse
import json
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    sys.exit("python-docx not installed. Run: pip install python-docx")


def set_cn_font(run, font_cn="宋体", font_en="Times New Roman", size=12):
    run.font.name = font_en
    run.font.size = Pt(size)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        from docx.oxml import OxmlElement
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), font_cn)


def add_heading(doc, text, level=1):
    style_map = {1: ("黑体", 16, True), 2: ("黑体", 14, True), 3: ("黑体", 12, True)}
    font_cn, size, bold = style_map.get(level, ("宋体", 12, False))
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_cn_font(r, font_cn=font_cn, size=size)
    r.bold = bold


def add_body(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_cn_font(r, size=12)
    p.paragraph_format.first_line_indent = Cm(0.74)


def add_table(doc, rows):
    if not rows:
        return
    tbl = doc.add_table(rows=len(rows), cols=len(rows[0]))
    tbl.style = "Light Grid Accent 1"
    for i, row in enumerate(rows):
        for j, cell in enumerate(row):
            c = tbl.cell(i, j)
            c.text = ""
            r = c.paragraphs[0].add_run(str(cell))
            set_cn_font(r, size=11)
            if i == 0:
                r.bold = True


def build(data: dict, out_path: Path):
    doc = Document()
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run(data["title"])
    set_cn_font(r, font_cn="黑体", size=18); r.bold = True
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run(data.get("subtitle", ""))
    set_cn_font(r, font_cn="黑体", size=14)
    # Lead
    if data.get("lead"):
        add_body(doc, data["lead"])

    for sec in data.get("sections", []):
        add_heading(doc, sec["heading"], level=1)
        for sub in sec.get("subsections", []):
            add_heading(doc, sub["heading"], level=2)
            for block in sub.get("blocks", []):
                if block["type"] == "paragraph":
                    add_body(doc, block["text"])
                elif block["type"] == "table":
                    add_table(doc, block["rows"])
                elif block["type"] == "subheading":
                    add_heading(doc, block["text"], level=3)

    doc.save(str(out_path))
    print(f"Wrote {out_path}")


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("json_path", type=Path)
    ap.add_argument("--out", type=Path, default=Path("peer_analysis.docx"))
    args = ap.parse_args()
    data = json.loads(args.json_path.read_text())
    build(data, args.out)


if __name__ == "__main__":
    main()
